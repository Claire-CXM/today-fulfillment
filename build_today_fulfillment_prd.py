from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.section import WD_ORIENT
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "今日兑现_产品需求方案_PRD.docx"
FLOW_IMAGE = ROOT / "today-fulfillment-flow.png"
ARCH_IMAGE = ROOT / "today-fulfillment-architecture.png"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "596273"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
GREEN = "CDF4D3"
YELLOW = "FFECBD"
RED = "FFCDC2"
VIOLET = "DCCCFF"


def set_font(run, name="Calibri", size=11, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for grid_col, width in zip(grid.gridCol_lst, widths):
        grid_col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width / 1440)
            tc_w = cell._tc.tcPr.tcW
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def add_table(doc, headers, rows, widths, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    hdr = table.rows[0]
    repeat_header(hdr)
    for cell, text in zip(hdr.cells, headers):
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(text)
        set_font(run, size=font_size, color=INK, bold=True)
    for row_values in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row_values):
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            run = p.add_run(str(value))
            set_font(run, size=font_size, color="202124")
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_body(doc, text, bold_lead=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.1
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        set_font(r, size=11, color=INK, bold=True)
        r = p.add_run(text[len(bold_lead):])
        set_font(r, size=11, color="202124")
    else:
        r = p.add_run(text)
        set_font(r, size=11, color="202124")
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.1
    r = p.add_run(text)
    set_font(r, size=10.5, color="202124")
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_font(r, size=9.2, color=MUTED, italic=True)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_font(run, size=9, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)
    run = paragraph.add_run(" 页")
    set_font(run, size=9, color=MUTED)


def set_header_footer(section):
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("今日兑现 | 产品需求文档")
    set_font(r, size=9, color=MUTED)
    footer = section.footer
    p = footer.paragraphs[0]
    add_page_number(p)


def create_architecture_image():
    width, height = 1800, 1320
    image = Image.new("RGB", (width, height), "#FFFFFF")
    draw = ImageDraw.Draw(image)
    font_path = "C:/Windows/Fonts/msyh.ttc"
    try:
        title_font = ImageFont.truetype(font_path, 46, index=0)
        module_font = ImageFont.truetype(font_path, 28, index=0)
        item_font = ImageFont.truetype(font_path, 21, index=0)
        small_font = ImageFont.truetype(font_path, 18, index=0)
    except OSError:
        title_font = module_font = item_font = small_font = ImageFont.load_default()

    def rounded(box, fill, outline, radius=22, width_line=4):
        draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width_line)

    def center_text(box, text, font, fill="#202124"):
        left, top, right, bottom = box
        bb = draw.multiline_textbbox((0, 0), text, font=font, spacing=5, align="center")
        tw, th = bb[2] - bb[0], bb[3] - bb[1]
        draw.multiline_text(((left + right - tw) / 2, (top + bottom - th) / 2), text, font=font, fill=fill, spacing=5, align="center")

    def arrow(x1, y1, x2, y2, fill="#72767A"):
        draw.line((x1, y1, x2, y2), fill=fill, width=4)
        if y2 > y1:
            draw.polygon([(x2, y2), (x2 - 10, y2 - 18), (x2 + 10, y2 - 18)], fill=fill)
        else:
            draw.polygon([(x2, y2), (x2 - 10, y2 + 18), (x2 + 10, y2 + 18)], fill=fill)

    draw.text((70, 50), "今日兑现 | 产品功能架构", font=title_font, fill="#0B2545")
    draw.line((70, 120, 1730, 120), fill="#D7E2EE", width=3)
    root = (700, 165, 1100, 260)
    rounded(root, "#FFC2EC", "#F849C1")
    center_text(root, "今日兑现", module_font, "#4A1741")

    modules = [
        ("任务与拆分", ["新增与编辑", "时长校验", "手动与 AI 节点"], "#C2E5FF", "#3DADFF"),
        ("专注执行", ["倒计时", "暂停与中断", "临期预警与申辩"], "#C6FAF6", "#5AD8CC"),
        ("激励与约束", ["文案风格", "随机与阶段奖励", "随机安全惩罚"], "#FFECBD", "#FFC943"),
        ("复盘与成长", ["绿黄红日历", "黄色任务补做", "日月总结与自由日"], "#CDF4D3", "#66D575"),
        ("账号与数据", ["手机验证码", "游客数据迁移", "导出与删除"], "#DCCCFF", "#874FFF"),
        ("提醒与设置", ["每日计划提醒", "减少动画", "隐私偏好"], "#E8EEF5", "#B3B3B3"),
    ]
    positions = [(95, 360), (655, 360), (1215, 360), (95, 850), (655, 850), (1215, 850)]
    for (name, items, fill, outline), (x, y) in zip(modules, positions):
        module_box = (x, y, x + 490, y + 82)
        arrow(900, 260, x + 245, y - 10)
        rounded(module_box, fill, outline)
        center_text(module_box, name, module_font, "#1E1E1E")
        for i, item in enumerate(items):
            item_y = y + 120 + i * 92
            item_box = (x + 35, item_y, x + 455, item_y + 62)
            rounded(item_box, "#FFFFFF", outline, radius=16, width_line=2)
            center_text(item_box, item, item_font)
            arrow(x + 245, y + 82, x + 245, item_y - 7)
    note = (95, 1230, 1705, 1292)
    rounded(note, "#F4F6F9", "#D7E2EE", radius=14, width_line=2)
    center_text(note, "初期形态：手机优先 PWA；后续再扩展原生 App 的应用使用监测与系统级约束能力。", small_font, "#596273")
    image.save(ARCH_IMAGE)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def build_document():
    create_architecture_image()
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(doc)
    set_header_footer(section)

    for _ in range(5):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("今日兑现")
    set_font(r, size=30, color=INK, bold=True)
    p.paragraph_format.space_after = Pt(6)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("产品需求文档（PRD）")
    set_font(r, size=18, color=BLUE, bold=True)
    p.paragraph_format.space_after = Pt(20)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("把今天的学习计划兑现")
    set_font(r, size=13, color=MUTED, italic=True)
    p.paragraph_format.space_after = Pt(34)
    add_table(doc, ["项目", "内容"], [
        ("产品形态", "手机优先的 PWA 学习助手"),
        ("版本", "V1.0 - 需求确认版"),
        ("文档日期", "2026 年 8 月 14 日"),
        ("当前状态", "产品方案确认，待进入开发"),
    ], [2600, 6760], font_size=10)
    doc.add_page_break()

    add_heading(doc, "1. 产品定位与目标", 1)
    add_body(doc, "今日兑现是一款面向课程学习、作业推进和转岗储备场景的学习承诺兑现助手。它不把用户带入复杂的待办管理，而是帮助用户每天聚焦最重要的三件事，并在计划、执行、反馈和复盘之间建立可持续的闭环。")
    add_heading(doc, "1.1 核心问题", 2)
    for text in [
        "计划模糊：学习任务通常过大，难以启动和衡量进度。",
        "过程失控：缺少时间边界、节点反馈和临期提醒，容易拖到最后。",
        "失败即放弃：一次中断或超时后容易放弃整天计划，缺少合理补做机制。",
        "正反馈不足：努力不可见，奖励和复盘缺位，难形成长期习惯。",
    ]:
        add_bullet(doc, text)
    add_heading(doc, "1.2 产品目标", 2)
    add_table(doc, ["目标", "可验证表现"], [
        ("提升启动率", "用户在提醒后能快速创建当天任务并开始第一项任务。"),
        ("提升完成率", "通过节点、倒计时和预警，使任务在时限内完成。"),
        ("降低挫败感", "特殊中断与黄色待补机制允许用户补回部分未完成任务。"),
        ("形成复盘习惯", "日总结、月总结和自由日卡让学习表现长期可见。"),
    ], [2600, 6760])

    add_heading(doc, "2. MVP 范围与边界", 1)
    add_table(doc, ["纳入初期版本", "暂不纳入初期版本"], [
        ("任务创建、节点拆分、倒计时、暂停和特殊中断", "跨应用使用时长监测"),
        ("免费 AI 拆分与规则模板兜底", "自动判断申辩真假"),
        ("随机奖励、随机安全惩罚、日历和复盘", "关机、卸载 App、账号登出等系统级操作"),
        ("手机号登录、游客试用、PWA 通知", "原生 App 专属的强制专注能力"),
    ], [4680, 4680])
    add_body(doc, "边界原则：初期产品以自我承诺、可解释的例外和正向反馈为主，不以不可逆的设备操作作为惩罚手段。", "边界原则：")

    doc.add_page_break()
    add_heading(doc, "3. 核心操作流程", 1)
    add_body(doc, "用户从创建当天任务进入专注执行。任务可被拆分为节点；系统在临期且明显落后时提醒；完成后进入奖励与复盘，超时后可申辩并转为黄色待补。")
    if FLOW_IMAGE.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(FLOW_IMAGE), width=Inches(4.55))
        add_caption(doc, "图 1  今日兑现核心操作流程（可编辑源文件已在 FigJam 保存）")
    else:
        add_body(doc, "图 1 导出文件缺失；请以 FigJam 中的核心流程图为准。")
    add_body(doc, "FigJam 可编辑源： https://www.figma.com/board/bqNBmZwh0NzhDzsfCpcP0z")

    doc.add_page_break()
    add_heading(doc, "4. 产品功能架构", 1)
    add_body(doc, "产品以任务与拆分、专注执行、激励与约束、复盘与成长、账号与数据、提醒与设置六大模块构成。模块间以任务状态和每日完成状态为主要数据连接点。")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(ARCH_IMAGE), width=Inches(6.15))
    add_caption(doc, "图 2  今日兑现产品功能架构（与 FigJam 架构图内容一致）")
    add_body(doc, "FigJam 可编辑源： https://www.figma.com/board/Ac4TwVvKKQD6o1Y5FhTgt1")

    doc.add_page_break()
    add_heading(doc, "5. 信息架构与关键页面", 1)
    add_table(doc, ["页面", "用户目的", "核心内容"], [
        ("登录与游客页", "开始使用并保留数据", "手机号验证码、游客试用、游客数据迁移"),
        ("今日主页", "安排并推进当天任务", "任务卡片、完成进度、今日专注时长、自由日状态"),
        ("任务编辑页", "定义可执行任务", "任务内容、预计时长、节点、AI 拆分、保存与删除"),
        ("专注页", "完成当前任务", "倒计时、节点勾选、暂停、特殊中断、完成任务"),
        ("日历与复盘页", "查看历史和补做", "绿黄红日历、遗留任务、日总结、月总结"),
        ("我的与设置", "管理激励和偏好", "奖励池、惩罚池、提醒、动效、导出与账号删除"),
    ], [1600, 2700, 5060])

    add_heading(doc, "6. 功能需求", 1)
    add_heading(doc, "6.1 计划与任务拆分", 2)
    add_table(doc, ["编号", "功能", "需求规则"], [
        ("FR-01", "每日任务", "新一天默认显示 3 个任务位，用户可随时增删；已开始任务删除需二次确认。"),
        ("FR-02", "时长校验", "保存新任务时，当前时间加预计时长不得晚于当日 24:00。"),
        ("FR-03", "单任务执行", "同一时间只有 1 个任务可进行；开始新任务前需暂停或完成当前任务。"),
        ("FR-04", "节点拆分", "每个任务最多 10 个节点；支持手动新增和 AI 生成，AI 结果必须可编辑确认。"),
        ("FR-05", "进度计算", "有节点任务按完成节点数除以总节点数计算；无节点任务支持手动调整或直接完成。"),
    ], [900, 1700, 6760])
    add_heading(doc, "6.2 专注、暂停与失败", 2)
    add_table(doc, ["编号", "功能", "需求规则"], [
        ("FR-06", "倒计时", "点击开始后按秒倒计时；完成时记录实际专注时长。"),
        ("FR-07", "普通暂停", "单任务最多暂停 3 次，累计最多 15 分钟；记录进入日总结。"),
        ("FR-08", "特殊中断", "用户填写原因后可无限时中断并保留进度；跨日后原日期转为黄色待补。"),
        ("FR-09", "临期预警", "剩余不超过 30 分钟且进度低于 60% 时，按用户选择的文案风格提醒一次。"),
        ("FR-10", "失败与申辩", "倒计时结束未完成即失败；仅可申辩一次，说明并诚信确认后转为黄色待补。"),
    ], [900, 1700, 6760])

    doc.add_page_break()
    add_heading(doc, "6.3 激励、约束与复盘", 2)
    add_table(doc, ["编号", "功能", "需求规则"], [
        ("FR-11", "完成反馈", "节点完成显示短激励；任务完成显示可关闭的烟花动效和鼓励弹窗。"),
        ("FR-12", "小奖励池", "用户维护 N 个奖励；任务完成随机显示 3 个，可换一批，最终只能领取 1 个。"),
        ("FR-13", "阶段奖励", "用户设定累计专注时长目标和大奖励；达到后解锁并记录。"),
        ("FR-14", "日历状态", "绿钩为全完成；黄横线为部分完成并可补做；红叉为零完成或全天无任务，不可补做。"),
        ("FR-15", "日总结", "当日全部任务完成时，生成完成情况、时长、暂停/申辩、优点、改进点和鼓励结语。"),
        ("FR-16", "自由日卡", "每累计 7 个绿钩日解锁 1 张；启用日无需创建任务，不触发未计划惩罚。"),
        ("FR-17", "每日提醒", "默认 10:00 提醒，可设置；30 分钟仍未创建任务，随机抽取预设安全惩罚。"),
        ("FR-18", "月总结", "每月 1 日生成上月完成情况、绿黄红分布、专注时长与改进建议。"),
        ("FR-19", "账号与数据", "提供手机号验证码登录、游客试用、学习记录导出及账号和云端数据删除。"),
    ], [900, 1700, 6760])

    add_heading(doc, "7. 业务规则与状态模型", 1)
    add_table(doc, ["对象", "状态", "关键规则"], [
        ("任务", "未开始、进行中、普通暂停、特殊中断、已完成、已失败、待补做", "特殊中断或申辩后的跨日未完成任务可进入待补做。"),
        ("日期", "绿钩、黄横线、红叉、自由日", "只有黄色日期可补完并转绿；自由日不计入绿黄红完成率。"),
        ("奖励", "待领取、已领取、阶段已解锁", "单任务只可领取一个小奖励，领取记录包含时间与对应任务。"),
        ("惩罚", "待执行、已确认", "系统仅随机抽取用户预设的安全项目，不执行系统级强制动作。"),
    ], [1300, 3000, 5060])
    add_heading(doc, "7.1 补充业务规则", 2)
    for text in [
        "已开始任务被删除时，保留为“主动放弃”记录，进入日总结而不从历史中抹除。",
        "当天已经变绿后若再次新增任务，日期回到进行中，任务完成后更新日总结。",
        "黄色任务保留原计划日期和实际补完日期，以支持月度复盘。",
        "免费 AI 模型不可用、限流或生成失败时，系统必须回退到可编辑的规则拆分模板。",
    ]:
        add_bullet(doc, text)

    doc.add_page_break()
    add_heading(doc, "8. 非功能与技术约束", 1)
    add_table(doc, ["维度", "要求"], [
        ("产品形态", "初期为手机优先、可安装的 PWA；后续可封装为小程序或原生 App。"),
        ("通知", "每日计划提醒依赖用户授权浏览器通知；不同设备的后台能力存在差异。"),
        ("AI", "采用可替换的免费模型适配层，核心任务创建流程不能依赖单一免费服务。"),
        ("短信", "正式手机号验证码需要接入短信服务；开发测试可使用测试号码或模拟验证码。"),
        ("隐私", "默认只收集任务功能必需数据；提供导出和删除；申辩内容按个人数据保护。"),
        ("后续增强", "应用使用时长检测、申辩核验及系统级限制仅在原生 App 阶段评估。"),
    ], [2000, 7360])

    add_heading(doc, "9. MVP 验收清单", 1)
    add_table(doc, ["场景", "验收结果"], [
        ("新增任务", "超过当日 24:00 的任务不能保存，页面给出调整建议。"),
        ("同时启动任务", "第二个任务不能直接开始，系统要求处理当前任务。"),
        ("节点完成", "勾选节点后总进度立即刷新并展示激励。"),
        ("临期落后", "剩余 30 分钟以内且进度低于 60% 时仅提醒一次。"),
        ("任务超时", "任务变为失败，申辩入口最多可使用一次。"),
        ("黄色补做", "补完遗留任务后，原日期从黄横线变为绿钩。"),
        ("奖励领取", "每次任务完成最多领取 1 个随机小奖励，并留下记录。"),
        ("每日提醒", "指定时间通知；30 分钟未计划触发随机安全惩罚。"),
        ("自由日", "启用自由日后当天免创建任务和未计划惩罚。"),
        ("数据管理", "游客数据可迁移；账号数据可导出和删除。"),
    ], [2500, 6860])

    add_heading(doc, "10. 开发优先级", 1)
    add_table(doc, ["阶段", "交付内容"], [
        ("P0 - 学习闭环", "今日任务、时长校验、倒计时、节点、暂停、完成、日历基本状态。"),
        ("P1 - 动力机制", "预警文案、奖励池、惩罚池、日总结、月总结、自由日卡。"),
        ("P2 - 智能与账号", "免费 AI 拆分与兜底模板、手机号登录、游客迁移、通知与数据管理。"),
        ("P3 - 原生增强", "应用使用时长检测、申辩核验、系统级限制能力的可行性评估。"),
    ], [2100, 7260])
    add_body(doc, "结论：今日兑现的 MVP 应优先验证“每天三件事是否更容易开始并完成”的学习闭环，再逐步增强 AI、账号和设备级约束能力。", "结论：")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_document()
